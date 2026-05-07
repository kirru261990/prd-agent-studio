from concurrent.futures import ThreadPoolExecutor
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import anthropic
from tavily import TavilyClient
from dotenv import load_dotenv
import os
import json
import re
from pathlib import Path
from datetime import datetime
import fitz

load_dotenv()

app = FastAPI()
client = anthropic.Anthropic()
tavily = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/ui", StaticFiles(directory=".", html=True), name="ui")

# RBI document registry
RBI_DOCS = {
    "Credit Card FAQs": "data/rbi/MD - Credit card FAQs.pdf",
    "Credit Cards Master Direction": "data/rbi/MD - Credit cards.pdf",
    "BBPS Guidelines": "data/rbi/MD - BBPS.pdf",
    "Credit Card Issuance Conduct 2022": "data/rbi/MD - Credit Card Issuance Conduct 2022.pdf"
}

# Input models
class ClarifyInput(BaseModel):
    feature_name: str
    feature_description: str
    competitors: str = ""
    qa_history: list = []

class AnalysisInput(BaseModel):
    feature_name: str
    feature_description: str
    competitors: str = ""
    qa_history: list = []


def extract_rbi_context(feature_description: str) -> str:
    keywords = [w for w in feature_description.lower().split()
                if len(w) > 3 and w not in
                ['this', 'that', 'with', 'from', 'they', 'their', 'will', 'have',
                 'been', 'into', 'than', 'then', 'when', 'what', 'hdfc', 'pixel',
                 'play', 'user', 'users', 'card', 'credit', 'feature']]

    all_relevant = []
    for doc_name, doc_path in RBI_DOCS.items():
        if not Path(doc_path).exists():
            continue
        doc = fitz.open(doc_path)
        full_text = "\n".join([page.get_text() for page in doc])
        lines = full_text.split('\n')
        relevant_lines = []
        for i, line in enumerate(lines):
            if len(line.strip()) < 5:
                continue
            if any(kw in line.lower() for kw in keywords):
                start = max(0, i - 2)
                end = min(len(lines), i + 6)
                chunk = "\n".join(lines[start:end])
                if chunk not in relevant_lines:
                    relevant_lines.append(chunk)
        if relevant_lines:
            all_relevant.append(
                f"--- {doc_name} ---\n" + "\n\n".join(relevant_lines[:8])
            )

    return "\n\n".join(all_relevant) if all_relevant else \
        "No specific RBI section found. Proceed with general compliance awareness."


def load_pixel_context() -> str:
    context_path = Path("pixel-context-private.md")
    return context_path.read_text() if context_path.exists() else ""


def save_analysis(feature_name: str, result: dict):
    output_dir = Path("data/analyses")
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r'[^a-z0-9]', '_', feature_name.lower())[:50]
    filename = output_dir / f"{timestamp}_{safe_name}.json"
    with open(filename, 'w') as f:
        json.dump(result, f, indent=2)
    return str(filename)


def format_qa_history(qa_history: list) -> str:
    if not qa_history:
        return ""
    lines = ["\nCLARIFYING QUESTIONS AND ANSWERS:"]
    for qa in qa_history:
        lines.append(f"Q: {qa['question']}")
        lines.append(f"A: {qa['answer']}")
    return "\n".join(lines)

def extract_options(response_text: str) -> dict:
    """Extract structured options JSON from agent response."""
    import re
    pattern = r'<options>(.*?)</options>'
    match = re.search(pattern, response_text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1).strip())
        except json.JSONDecodeError:
            return {}
    return {}

def extract_structured_options(response_text: str, agent_type: str) -> dict:
    """Use a separate Claude call to extract structured options from agent response."""
    
    if agent_type == "competitor":
        extraction_prompt = f"""Extract the 3 differentiation hypotheses from this competitive analysis and return ONLY a JSON object with no other text.

Analysis:
{response_text[:3000]}

Return exactly this JSON structure:
{{"section": "differentiation", "agent": "competitor", "choices": [{{"id": "angle_1", "title": "3-4 word title", "summary": "1 sentence what to build", "rationale": "1 sentence why", "complexity": "Low/Medium/High", "recommended": true}}, {{"id": "angle_2", "title": "3-4 word title", "summary": "1 sentence what to build", "rationale": "1 sentence why", "complexity": "Low/Medium/High", "recommended": false}}, {{"id": "angle_3", "title": "3-4 word title", "summary": "1 sentence what to build", "rationale": "1 sentence why", "complexity": "Low/Medium/High", "recommended": false}}]}}"""

    else:  # cs agent
        extraction_prompt = f"""Extract the top 3 ticket deflection opportunities from this CS analysis and return ONLY a JSON object with no other text.

Analysis:
{response_text[:3000]}

Return exactly this JSON structure:
{{"section": "cs_strategy", "agent": "cs", "choices": [{{"id": "cs_1", "title": "3-4 word title", "summary": "1 sentence what to build", "rationale": "1 sentence why this reduces tickets", "ticket_reduction": "estimated % e.g. 40%", "complexity": "Low/Medium/High", "recommended": true}}, {{"id": "cs_2", "title": "3-4 word title", "summary": "1 sentence", "rationale": "1 sentence", "ticket_reduction": "estimated %", "complexity": "Low/Medium/High", "recommended": false}}, {{"id": "cs_3", "title": "3-4 word title", "summary": "1 sentence", "rationale": "1 sentence", "ticket_reduction": "estimated %", "complexity": "Low/Medium/High", "recommended": false}}]}}"""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=1000,
            temperature=0.0,
            messages=[{"role": "user", "content": extraction_prompt}]
        )
        response = message.content[0].text.strip()
        # Clean any markdown code blocks if present
        response = re.sub(r'```json\s*|\s*```', '', response).strip()
        return json.loads(response)
    except Exception as e:
        return {}


@app.get("/health")
def health():
    return {"status": "running", "agent": "prd-agent-studio"}


@app.post("/clarify")
def clarify(input: ClarifyInput):
    """Generate next clarifying question or signal READY if enough context."""

    pixel_context = load_pixel_context()
    qa_so_far = format_qa_history(input.qa_history)
    question_number = len(input.qa_history) + 1

    prompt = f"""You are a senior product analyst. A PM has described a feature and you need to ask ONE clarifying question before running competitive research.

    Feature name: {input.feature_name}
    Feature description: {input.feature_description}
    {qa_so_far}

    Study the feature description carefully. Identify the single most important thing that is ambiguous, missing, or unclear that would most change the direction of the competitive research if you knew it.

    Do NOT ask about:
    - Things already clearly stated in the description
    - Generic questions like "who is the target user" if that's already obvious
    - Regulatory context if the feature already mentions RBI
    - Things that won't change the research direction

    DO ask about:
    - The primary business goal if it's ambiguous (cost saving vs revenue vs retention vs compliance)
    - Edge cases that would change the competitive set entirely
    - Constraints that would eliminate certain differentiation angles
    - The specific trigger or moment in the user journey if unclear

    If after {len(input.qa_history)} answers you have enough context for excellent research, respond with: READY

    Otherwise respond with ONE sharp, specific question. No preamble. No explanation. Just the question."""

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=200,
        temperature= 0.0,
        system=f"You are a senior product analyst for HDFC PIXEL Studio.\n\n{pixel_context}",
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = message.content[0].text.strip()
    is_ready = response_text.upper() == "READY" or question_number > 3

    return {
        "question": response_text if not is_ready else None,
        "ready": is_ready,
        "question_number": question_number
    }


@app.post("/competitor-agent")
def competitor_agent(input: AnalysisInput):

    pixel_context = load_pixel_context()

    # Build enriched feature description from Q&A
    qa_context = format_qa_history(input.qa_history)
    enriched_description = f"{input.feature_description}\n{qa_context}"

    feature_short = input.feature_name[:100]

    # Detect RBI feature
    is_rbi_feature = any(word in enriched_description.lower()
                         for word in ['rbi', 'regulatory', 'compliance',
                                      'guideline', 'circular', 'mandate',
                                      'refund', 'bbps', 'autopay'])

    # Build competitor search list
    default_competitors = "OneCard Scapia Slice Kreditbee Freo Uni"
    custom_competitors = input.competitors.strip() if input.competitors.strip() else ""
    all_competitors = f"{custom_competitors} {default_competitors}".strip()

    # Step 1a: Broad discovery
    discovery_results = tavily.search(
        query=f"India fintech credit card {feature_short} 2024",
        max_results=3,
        search_depth="basic",
        include_raw_content=True,
        days=365,
    )

    # Step 1b: Individual searches per top competitor
    top_competitors = ["OneCard", "Scapia", "Slice"]
    if custom_competitors:
        top_competitors = [c.strip() for c in custom_competitors.split(",")][:3]

    competitor_results_list = []
    for comp in top_competitors:
        try:
            result = tavily.search(
                query=f"{comp} {feature_short} India feature review",
                max_results=2,
                search_depth="advanced",
                include_raw_content=True,
                days=365,
            )
            competitor_results_list.extend(result.get('results', []))
        except Exception:
            continue

    competitor_results = {"results": competitor_results_list}

    # Step 1c: Legacy bank search
    legacy_results = tavily.search(
        query=f"ICICI Axis Kotak SBI {feature_short} credit card India feature",
        max_results=3,
        search_depth="basic",
        include_raw_content=True,
        days=365,
    )

    # Step 1d: Community signal — Reddit and Quora
    community_results = tavily.search(
        query=f"reddit OR quora {feature_short} credit card India review complaint experience 2024",
        max_results=3,
        search_depth="basic",
        include_raw_content=True,
        days=365,
    )

    # Step 1e: YouTube signal
    youtube_results = tavily.search(
        query=f"youtube {feature_short} credit card India review explained 2024",
        max_results=2,
        search_depth="basic",
        include_raw_content=True,
        days=365,
    )
    

   # Deduplicate all results
    all_results = (
        discovery_results.get('results', []) +
        competitor_results.get('results', []) +
        legacy_results.get('results', []) +
        community_results.get('results', []) +
        youtube_results.get('results', [])
    )
    
    seen_urls = set()
    unique_results = []
    for r in all_results:
        if r['url'] not in seen_urls:
            seen_urls.add(r['url'])
            unique_results.append(r)

    def is_relevant(result):
        url = result.get('url', '').lower()
        content = result.get('content', '').lower()
        # Filter out generic marketing pages and irrelevant domains
        skip_domains = ['wikipedia', 'investopedia', 'bankbazaar', 
                       'policybazaar', 'moneycontrol']
        if any(d in url for d in skip_domains):
            return False
        # Must contain feature keywords
        feature_words = feature_short.lower().split()[:3]
        return any(w in content for w in feature_words)

    filtered_results = [r for r in unique_results if is_relevant(r)]
    if not filtered_results:
        filtered_results = unique_results  # fallback if all filtered out

    competitor_context = "\n\n".join([
       f"Source: {r['url']}\n{(r.get('raw_content') or r.get('content') or '')[:2000]}"
        for r in filtered_results[:8]
    ])

    # Step 2: RBI or global context
    if is_rbi_feature:
        secondary_context = extract_rbi_context(enriched_description)
        secondary_label = "RBI REGULATORY CONTEXT"
    else:
        global_results = tavily.search(
            query=f"global best practice {feature_short} credit card fintech Monzo Revolut 2024",
            max_results=3,
            search_depth="basic",
            include_raw_content=True,
            days=365,
        )
        secondary_context = "\n\n".join([
            f"Source: {r['url']}\n{r['content']}"
            for r in global_results.get('results', [])
        ])
        secondary_label = "GLOBAL BEST PRACTICES"

    # Step 3: Build prompt
    prompt = f"""A PM has asked for competitive intelligence on this feature:

FEATURE NAME: {input.feature_name}
FEATURE DESCRIPTION: {input.feature_description}
{qa_context}

IMPORTANT RULES:
- HDFC Bank, HDFC PIXEL Play, and PayZapp are the products we are BUILDING — never list them as competitors
- Use ONLY the search results below — do not use training memory
- If something is not in sources, say "not found in current sources"
- Competitor priority: Digital-first cards (OneCard, Scapia, Slice) → Legacy banks (ICICI, Axis, Kotak) → NBFCs/Fintechs (KreditBee, Freo)

SEARCH RESULTS:
{competitor_context}

{secondary_label}:
{secondary_context}

Respond in exactly this structure:

COMPETITORS:
List up to 5. Tag each with priority: [DIGITAL-FIRST] [LEGACY BANK] or [FINTECH/NBFC]
For each:
- Name | Product name [PRIORITY TAG]
- Why relevant: (1 sentence)
- What they built: (2 sentences — specific, not generic)
- One key point: (1 sentence)
- Confidence: High (has URL) / Medium (inferred) / Low (assumed)
- Source: (URL or "not found")

WHITESPACE SIGNAL:
If fewer than 2 competitors found with this specific feature:
"FIRST-MOVER OPPORTUNITY: [explain the gap in 1 sentence]"
Otherwise skip.

WHAT REAL USERS SAY:
2-3 observations from Reddit or app store reviews with source URLs.
If none found: "No community signal found in current sources."

{secondary_label} SUMMARY:
If RBI: Quote exact RBI rule. List what PIXEL must comply with as bullet points.
If global: 2-3 examples, what works, why, source URL.

DIFFERENTIATION HYPOTHESES:
Generate exactly 3 differentiation angles for PIXEL. 

CRITICAL RULE: Each angle must represent a fundamentally different strategic choice — different positioning, different target user segment, different build complexity, or different competitive moat. A PM should feel genuine tension choosing between them. They must NOT be sequential steps of the same feature.

Think across these dimensions for contrast:
- Simple vs sophisticated (zero-UI automation vs intelligent personalisation)
- Defensive vs offensive (compliance-minimum vs market leadership)
- Transactional vs relational (solve the immediate problem vs create a deeper relationship moment)
- Fast to ship vs high moat (quick win vs durable advantage)

For each angle:
- Angle name: (3-4 words)
- What to build: (1 sentence — specific, not generic)
- Strategic logic: (1 sentence — why this is a distinct bet)
- RBI compliance: (1 sentence)
- Build complexity: Low / Medium / High
- Ops impact: (1 sentence)

RECOMMENDED ANGLE:
Pick the strongest angle for PIXEL's current stage and constraints.
- Which angle: (name it)
- Why now: (why this angle fits PIXEL's current technical maturity and team capacity)
- Why not the others: (1 sentence each explaining the tradeoff)

OVERALL CONFIDENCE: High / Medium / Low
Reason: (1 sentence)"""

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=3000,
        temperature= 0.0,
        system=[
            {
                "type": "text",
                "text": f"You are a senior competitive intelligence analyst for HDFC PIXEL Studio. HDFC Bank, PIXEL Play, and PayZapp are the products being built — never list them as competitors.\n\nPIXEL Studio context:\n{pixel_context}",
                "cache_control": {"type": "ephemeral"}
            }
        ],
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = message.content[0].text
    options = extract_structured_options(response_text, "competitor")

    result = {
        "agent": "competitor",
        "feature_name": input.feature_name,
        "timestamp": datetime.now().isoformat(),
        "sources_searched": len(unique_results),
        "rbi_triggered": is_rbi_feature,
        "qa_history": input.qa_history,
        "response": response_text,
        "options": options
    }

    # Save analysis
    saved_path = save_analysis(input.feature_name, result)
    result["saved_to"] = saved_path

    return result
    

# ============================================================
# CS AGENT
# ============================================================

class CSInput(BaseModel):
    feature_name: str
    feature_description: str
    competitors: str = ""
    qa_history: list = []


def load_prd_context(feature_name: str) -> str:
    """Find and extract most relevant PRD from /data/prds/ based on feature name."""
    from docx import Document

    prds_dir = Path("data/prds")
    if not prds_dir.exists():
        return "No PRDs found in data/prds directory."

    # Find all docx files
    prd_files = list(prds_dir.glob("*.docx"))
    if not prd_files:
        return "No PRD files found."

    # Score each PRD by keyword match with feature name
    feature_keywords = [w.lower() for w in feature_name.split()
                        if len(w) > 3]

    best_file = None
    best_score = 0

    for prd_file in prd_files:
        filename_lower = prd_file.name.lower()
        score = sum(1 for kw in feature_keywords if kw in filename_lower)
        if score > best_score:
            best_score = score
            best_file = prd_file

    # If no keyword match, use first PRD
    if not best_file:
        best_file = prd_files[0]

    # Extract text from docx
    try:
        doc = Document(best_file)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return f"PRD: {best_file.name}\n\n{text[:4000]}"
    except Exception as e:
        return f"Could not read PRD: {str(e)}"


@app.post("/cs-agent")
def cs_agent(input: CSInput):

    pixel_context = load_pixel_context()
    qa_context = format_qa_history(input.qa_history)
    enriched_description = f"{input.feature_description}\n{qa_context}"
    feature_short = input.feature_name[:100]

    # Load relevant PRD
    prd_context = load_prd_context(input.feature_name)

    # Detect RBI feature
    is_rbi_feature = any(word in enriched_description.lower()
                         for word in ['rbi', 'regulatory', 'compliance',
                                      'guideline', 'circular', 'mandate',
                                      'refund', 'bbps', 'autopay'])

    # Search 1: Competitor CX handling
    cx_results = tavily.search(
        query=f"OneCard Scapia Slice {feature_short} customer support FAQ self-serve India",
        max_results=3,
        search_depth="basic",
        include_raw_content=True,
        days=365,
    )

    # Search 2: Community signal — what users complain about
    complaint_results = tavily.search(
        query=f"reddit quora {feature_short} credit card India problem complaint support 2024",
        max_results=3,
        search_depth="basic",
        include_raw_content=True,
        days=365,
    )

    # Search 3: Global CX best practice
    global_cx_results = tavily.search(
        query=f"best practice {feature_short} credit card customer support deflection self-serve",
        max_results=2,
        search_depth="basic",
        include_raw_content=True,
        days=365,
    )

    # Format search results
    cx_context = "\n\n".join([
        f"Source: {r['url']}\n{r['content']}"
        for r in cx_results.get('results', [])
    ])

    complaint_context = "\n\n".join([
        f"Source: {r['url']}\n{r['content']}"
        for r in complaint_results.get('results', [])
    ])

    global_cx_context = "\n\n".join([
        f"Source: {r['url']}\n{r['content']}"
        for r in global_cx_results.get('results', [])
    ])

    # RBI context if applicable
    rbi_context = ""
    if is_rbi_feature:
        rbi_context = f"\nRBI REGULATORY CONTEXT:\n{extract_rbi_context(enriched_description)}"

    # Build prompt
    prompt = f"""You are a senior customer support strategist reviewing a PRD for HDFC PIXEL Studio.

FEATURE NAME: {input.feature_name}
FEATURE DESCRIPTION: {input.feature_description}
{qa_context}

PRD CONTENT:
{prd_context}

{rbi_context}

COMPETITOR CX SEARCH RESULTS:
{cx_context}

USER COMPLAINT SIGNAL:
{complaint_context}

GLOBAL CX BEST PRACTICES:
{global_cx_context}

Your job is to identify every CS impact this feature will have and how to handle it.

IMPORTANT RULES:
- Be specific to this feature — no generic CS advice
- Reference the PRD content where relevant
- Flag gaps in the PRD that will cause CS problems
- Use search results to benchmark competitor CX approaches

Respond in exactly this structure:

CONTACT DRIVERS (maximum 5, highest likelihood only):
List the 5 most likely reasons a user will contact CS.
For each:
- Scenario: (what happened from user's perspective)
- Likelihood: High / Medium / Low
- PRD gap: (is this covered in the PRD? Yes / Partial / No)

COMPETITOR CX BENCHMARK:
For each high-likelihood contact driver:
- How does the best competitor handle this? (specific, not generic)
- Self-serve or agent-assisted?
- What PIXEL should copy or improve on
- Source URL

FAQ COVERAGE (maximum 5 only):
For each contact driver write one FAQ:
- Question: (exactly how a user would phrase it)
- Answer: (what the CS agent should say — specific, actionable)
- Escalation needed: Yes / No

AGENT CENTRE REQUIREMENTS:
What information must be visible to CS agents to resolve queries?
List as bullet points. Flag any that are NOT currently in the PRD.

ERROR SCREEN AUDIT:
List every error state this feature introduces.
For each:
- Error scenario
- Current PRD handling: (documented / not documented)
- Recommended next action for user

ESCALATION PATHS:
Which scenarios cannot be resolved by tier-1 CS?
For each:
- Scenario
- Who handles it (tier-2 / bank team / Zeta ops)
- Expected resolution time

TICKET DEFLECTION OPPORTUNITIES:
What self-serve features would eliminate CS contact entirely?
For each opportunity:
- What to build
- Estimated ticket reduction %
- Complexity: Low / Medium / High

OVERALL CS READINESS SCORE: X/10
Reason: (2 sentences — what's good, what's missing)"""

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=3000,
        temperature= 0.0,
        system=[
            {
                "type": "text",
                "text": f"You are a senior customer support strategist for HDFC PIXEL Studio. You review PRDs to identify CS impact, gaps, and deflection opportunities.\n\nPIXEL Studio context:\n{pixel_context}",
                "cache_control": {"type": "ephemeral"}
            }
        ],
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = message.content[0].text
    options = extract_structured_options(response_text, "cs")

    result = {
        "agent": "cs",
        "feature_name": input.feature_name,
        "timestamp": datetime.now().isoformat(),
        "prd_used": prd_context.split('\n')[0],
        "rbi_triggered": is_rbi_feature,
        "qa_history": input.qa_history,
        "response": response_text,
        "options": options
    }

    saved_path = save_analysis(input.feature_name + "_cs", result)
    result["saved_to"] = saved_path

    return result

# ============================================================
# UNIFIED PRD REVIEW — runs competitor + CS agents in parallel
# ============================================================

class ReviewInput(BaseModel):
    feature_name: str
    feature_description: str
    competitors: str = ""
    qa_history: list = []


@app.post("/review")
def unified_review(input: ReviewInput):

    review_input_dict = {
        "feature_name": input.feature_name,
        "feature_description": input.feature_description,
        "competitors": input.competitors,
        "qa_history": input.qa_history
    }

    def run_competitor():
        return competitor_agent(AnalysisInput(**review_input_dict))

    def run_cs():
        return cs_agent(CSInput(**review_input_dict))

    def run_compliance():
        return compliance_agent(ComplianceInput(**review_input_dict))

    def run_spec():
        return spec_agent(SpecInput(**review_input_dict))

    # Fire all four agents in parallel
    with ThreadPoolExecutor(max_workers=4) as executor:
        competitor_future = executor.submit(run_competitor)
        cs_future = executor.submit(run_cs)
        compliance_future = executor.submit(run_compliance)
        spec_future = executor.submit(run_spec)

        competitor_result = competitor_future.result()
        cs_result = cs_future.result()
        compliance_result = compliance_future.result()
        spec_result = spec_future.result()

    # Combine results
    combined = {
        "agent": "unified_review",
        "feature_name": input.feature_name,
        "timestamp": datetime.now().isoformat(),
        "competitor_analysis": competitor_result,
        "cs_analysis": cs_result,
        "compliance_analysis": compliance_result,
        "spec_analysis": spec_result,
        "sources_searched": competitor_result.get("sources_searched", 0),
        "rbi_triggered": competitor_result.get("rbi_triggered", False),
        "qa_history": input.qa_history
    }

    # Save combined result
    saved_path = save_analysis(input.feature_name + "_review", combined)
    combined["saved_to"] = saved_path

    return combined


# ============================================================
# PRD GENERATOR
# ============================================================

class PRDGenerateInput(BaseModel):
    feature_name: str
    feature_description: str
    selected_options: dict = {}
    qa_history: list = []


@app.post("/generate-prd")
def generate_prd(input: PRDGenerateInput):

    pixel_context = load_pixel_context()
    qa_context = format_qa_history(input.qa_history)

    # Extract selected options
    diff_option = input.selected_options.get('differentiation', {})
    cs_option = input.selected_options.get('cs_strategy', {})

    diff_text = f"{diff_option.get('title', '')}: {diff_option.get('summary', '')}" \
        if diff_option else "Not selected"
    cs_text = f"{cs_option.get('title', '')}: {cs_option.get('summary', '')}" \
        if cs_option else "Not selected"

    prompt = f"""You are an Associate Director of Products at a fintech company building credit card products for Indian banks. You are writing a PRD that will be reviewed by your VP of Products, who is senior, time-pressed, and will challenge any vague claim, made-up metric, or missing dependency.

The PM has provided the following real context. Use this as the foundation — do not make up replacements for anything provided here:

PRODUCT IDEA: {input.feature_name}

FEATURE DESCRIPTION: {input.feature_description}

{qa_context}

CHOSEN DIFFERENTIATION ANGLE (from competitive research):
{diff_text}

CHOSEN CS STRATEGY (from CS impact analysis):
{cs_text}

PIXEL STUDIO CONTEXT:
{pixel_context}

Generate a PRD in this exact sequence:

1. Strategic Context (Why Now)
- Based strictly on the feature description and business trigger provided above
- What is the opportunity cost of not building this in the next two quarters?
- 3 sentences maximum. No fluff.

2. Persona Definition
- Define 2-3 distinct user segments based on PIXEL Studio's customer profile
- For each persona state:
  a. Who they are (specific — e.g. "salaried professional, 28-35, Tier 2 city, spends ₹20-40k/month on card")
  b. Their primary behaviour today
  c. Their specific pain point
  d. What would make them adopt this feature?

3. Problem Statement
- Do NOT mention the solution or product name
- State the user pain, the current behaviour, and the business consequence
- 3 sentences maximum

4. Current State Metrics (Baseline)
- State what metrics establish the baseline
- If any baseline is unknown, state "unknown — measurement needed" and name the team that owns that data
- Do not propose targets here — this section is diagnosis only

5. Scope
- In scope: features only, no adjectives
- Out of scope: at least 2 items a stakeholder might reasonably expect, each with a one-line reason for exclusion

6. User Stories with Acceptance Criteria
- 3-5 stories, each mapped to a named persona from section 2
- Format: "As a [persona name], I want [specific action] so that [measurable outcome]"
- No vague words — replace "seamless", "easy", "personalized" with specific observable behaviour
- Each story must have:
  a. 2-3 acceptance criteria in "Given [context], when [action], then [result]" format
  b. One failure-state AC: what the user sees when something goes wrong

7. Success Metrics
- Split into two views:
  a. What the user gets: engagement, satisfaction, behaviour change
  b. What the bank gets: revenue impact, cost reduction, risk improvement
- For each metric: name, current baseline, target, timeframe, owner

8. Release Plan
- 3 phases: MVP, V1, V2
- MVP must be embarrassingly small — one core job done well
- For each phase:
  a. What ships
  b. Rollout: % of users first, what triggers expansion
  c. Kill switch condition
  d. Success gate before next phase
  e. Realistic timeline
  f. Compliance sign-off required before this phase ships

9. Operations & Customer Support
Based on chosen CS strategy: {cs_text}
- New support ticket types this will generate
- CS training and tooling needs before launch
- SLAs for new issue types
- Escalation path when dependency fails
- Fraud or abuse surface created, and mitigation

10. Reports & Reconciliation
- What new reports does this feature require? Who consumes them and at what frequency?
- What reconciliation processes are needed?
- What happens when reconciliation breaks?
- Any data extract requirements for downstream teams?

11. Open Questions
- Exactly 3 questions
- At least one strategic, one technical dependency
- Format: Question → Why it matters → Who owns the answer

Be specific to Indian fintech, UPI, and RBI regulatory context. No generic platitudes."""

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        temperature=0.0,
        system=f"You are an Associate Director of Products at HDFC PIXEL Studio writing a production-grade PRD.\n\n{pixel_context}",
        messages=[{"role": "user", "content": prompt}]
    )

    result = {
        "agent": "prd_generator",
        "feature_name": input.feature_name,
        "timestamp": datetime.now().isoformat(),
        "selected_options": input.selected_options,
        "prd": message.content[0].text
    }

    saved_path = save_analysis(input.feature_name + "_prd", result)
    result["saved_to"] = saved_path

    return result


# ============================================================
# COMPLIANCE AGENT
# ============================================================

class ComplianceInput(BaseModel):
    feature_name: str
    feature_description: str
    qa_history: list = []


def load_all_rbi_docs() -> str:
    """Load all RBI PDFs and return combined text."""
    all_text = []
    for doc_name, doc_path in RBI_DOCS.items():
        if not Path(doc_path).exists():
            continue
        doc = fitz.open(doc_path)
        text = "\n".join([page.get_text() for page in doc])
        all_text.append(f"=== {doc_name} ===\n{text}")
    return "\n\n".join(all_text)


def extract_relevant_rbi_sections(feature_description: str, full_rbi_text: str) -> str:
    """Extract most relevant RBI sections using keyword matching."""
    keywords = [w.lower() for w in feature_description.split()
                if len(w) > 3 and w.lower() not in
                ['this', 'that', 'with', 'from', 'they', 'their',
                 'will', 'have', 'been', 'into', 'than', 'then',
                 'when', 'what', 'hdfc', 'pixel', 'user', 'users']]

    lines = full_rbi_text.split('\n')
    relevant_chunks = []
    seen = set()

    for i, line in enumerate(lines):
        if len(line.strip()) < 10:
            continue
        if any(kw in line.lower() for kw in keywords):
            start = max(0, i - 3)
            end = min(len(lines), i + 8)
            chunk = "\n".join(lines[start:end])
            if chunk not in seen:
                seen.add(chunk)
                relevant_chunks.append(chunk)

    return "\n\n---\n\n".join(relevant_chunks[:15]) if relevant_chunks else \
        "No specific RBI section found for this feature."


@app.post("/compliance-agent")
def compliance_agent(input: ComplianceInput):

    pixel_context = load_pixel_context()
    qa_context = format_qa_history(input.qa_history)
    enriched_description = f"{input.feature_description}\n{qa_context}"

    # Load all RBI docs
    full_rbi_text = load_all_rbi_docs()
    relevant_rbi = extract_relevant_rbi_sections(enriched_description, full_rbi_text)

    # Search for additional regulatory context
    feature_short = input.feature_name[:80]
    reg_results = tavily.search(
        query=f"RBI guidelines {feature_short} credit card India compliance 2024",
        max_results=3,
        search_depth="basic",
        include_raw_content=True,
        days=365
    )

    reg_context = "\n\n".join([
        f"Source: {r['url']}\n{(r.get('raw_content') or r.get('content') or '')[:1500]}"
        for r in reg_results.get('results', [])
    ])

    prompt = f"""You are a senior regulatory compliance analyst for HDFC PIXEL Studio credit card products in India.

FEATURE: {input.feature_name}
DESCRIPTION: {input.feature_description}
{qa_context}

Your job is to identify every RBI regulatory requirement relevant to this feature and produce a compliance checklist that the PM can use directly in the PRD.

RBI MASTER DIRECTIONS AND CIRCULARS (relevant sections extracted):
{relevant_rbi}

ADDITIONAL REGULATORY SEARCH RESULTS:
{reg_context}

CRITICAL RULES:
- Quote RBI text VERBATIM where found — use exact words, cite the document and clause
- Never paraphrase RBI requirements — exact quotes only
- If a requirement is NOT found in the documents, say "Not found in available circulars — legal review recommended"
- Distinguish between MANDATORY requirements (shall/must) and ADVISORY (should/may)
- Flag any compliance gaps that would block launch

Respond in exactly this structure:

APPLICABLE RBI REGULATIONS:
For each relevant regulation found:
- Clause reference: (document name + clause number if available)
- Verbatim quote: (exact text from RBI document in quotes)
- What it means for this feature: (1 sentence plain English interpretation)
- Mandatory or Advisory: (Mandatory / Advisory)

COMPLIANCE REQUIREMENTS FOR THIS FEATURE:
List every specific thing PIXEL must build or ensure to comply:
- Requirement: (what must be done)
- Source: (which RBI clause)
- Currently in PRD: (Yes / No / Unknown)
- Risk if ignored: (Low / Medium / High — 1 sentence on consequence)

COMPLIANCE GAPS (maximum 5, most critical only):
Top 5 gaps not covered by available RBI documents:
- Gap: (what's unclear)
- Recommendation: (get legal opinion / check HDFC compliance team / low risk)

LAUNCH BLOCKERS:
List any compliance requirement that MUST be met before this feature can go live.
If none: state "No mandatory compliance blockers identified in available documents."

OVERALL COMPLIANCE RISK: Low / Medium / High
Reason: (2 sentences)"""

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=3000,
        temperature=0.0,
        system=[
            {
                "type": "text",
                "text": f"You are a senior regulatory compliance analyst for HDFC PIXEL Studio. You quote RBI regulations verbatim and never paraphrase regulatory requirements.\n\nPIXEL context:\n{pixel_context}",
                "cache_control": {"type": "ephemeral"}
            }
        ],
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = message.content[0].text

    result = {
        "agent": "compliance",
        "feature_name": input.feature_name,
        "timestamp": datetime.now().isoformat(),
        "rbi_sections_found": relevant_rbi != "No specific RBI section found for this feature.",
        "response": response_text
    }

    saved_path = save_analysis(input.feature_name + "_compliance", result)
    result["saved_to"] = saved_path

    return result


# ============================================================
# SPEC AGENT — Use cases, acceptance criteria, metrics, instrumentation
# ============================================================

class SpecInput(BaseModel):
    feature_name: str
    feature_description: str
    qa_history: list = []


def load_relevant_sop(feature_name: str) -> str:
    """Find and load the most relevant SOP from ops-sops folder."""
    sops_dir = Path("data/internal/ops-sops")
    if not sops_dir.exists():
        return "No SOPs found."

    feature_keywords = [w.lower() for w in feature_name.split()
                   if len(w) > 3 and w.lower() not in
                   ['pixel', 'play', 'credit', 'card', 'users', 'hdfc']]

    sop_files = list(sops_dir.glob("*.docx")) + \
                list(sops_dir.glob("*.doc"))

    best_file = None
    best_score = 0

    for sop_file in sop_files:
        filename_lower = sop_file.name.lower()
        score = sum(1 for kw in feature_keywords if kw in filename_lower)
        if score > best_score:
            best_score = score
            best_file = sop_file

    if not best_file and sop_files:
        # Try content-based matching on first 500 chars
        for sop_file in sop_files:
            try:
                if sop_file.suffix == '.docx':
                    from docx import Document
                    doc = Document(sop_file)
                    text = " ".join([p.text for p in doc.paragraphs
                                    if p.text.strip()])[:500].lower()
                    score = sum(1 for kw in feature_keywords if kw in text)
                    if score > best_score:
                        best_score = score
                        best_file = sop_file
            except Exception:
                continue

    if not best_file:
        return "No relevant SOP found."

    try:
        if best_file.suffix == '.docx':
            from docx import Document
            doc = Document(best_file)
            text = "\n".join([p.text for p in doc.paragraphs
                             if p.text.strip()])
            return f"SOP: {best_file.name}\n\n{text[:5000]}"
        else:
            return f"SOP found: {best_file.name} (format not supported for extraction)"
    except Exception as e:
        return f"Could not read SOP: {str(e)}"


def load_instrumentation_reference() -> str:
    """Load instrumentation reference from xlsx."""
    import openpyxl
    inst_path = Path("data/internal/instrumentation-reference.xlsx")
    if not inst_path.exists():
        return "Instrumentation reference not found."

    try:
        wb = openpyxl.load_workbook(inst_path)
        all_text = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            sheet_text = [f"Sheet: {sheet}"]
            for row in ws.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    sheet_text.append(" | ".join(
                        str(c) for c in row if c is not None
                    ))
            all_text.append("\n".join(sheet_text))
        return "\n\n".join(all_text)[:4000]
    except Exception as e:
        return f"Could not read instrumentation reference: {str(e)}"


@app.post("/spec-agent")
def spec_agent(input: SpecInput):

    pixel_context = load_pixel_context()
    qa_context = format_qa_history(input.qa_history)
    enriched_description = f"{input.feature_description}\n{qa_context}"
    feature_short = input.feature_name[:80]

    # Load relevant SOP
    sop_context = load_relevant_sop(input.feature_name)

    # Load instrumentation reference only if feature has UI/UX
    has_ui = any(word in enriched_description.lower() for word in
                 ['app', 'screen', 'button', 'ui', 'ux', 'user interface',
                  'flow', 'page', 'view', 'tap', 'click', 'display', 'show'])
    instrumentation_context = load_instrumentation_reference() if has_ui else \
        "Not applicable — no UI/UX component detected for this feature."

    # Load most relevant PRD
    prd_context = load_prd_context(input.feature_name)

    # Search for use case patterns
    uc_results = tavily.search(
        query=f"{feature_short} credit card India use cases edge cases 2024",
        max_results=3,
        search_depth="basic",
        include_raw_content=True,
        days=365
    )

    uc_context = "\n\n".join([
        f"Source: {r['url']}\n{(r.get('raw_content') or r.get('content') or '')[:1500]}"
        for r in uc_results.get('results', [])
    ])

    prompt = f"""You are a senior product analyst writing the specification section of a PRD for HDFC PIXEL Studio.

FEATURE: {input.feature_name}
DESCRIPTION: {input.feature_description}
{qa_context}

CURRENT STATE (from ops SOP):
{sop_context}

EXISTING PRD REFERENCE:
{prd_context}

INSTRUMENTATION REFERENCE:
{instrumentation_context}

EXTERNAL USE CASE RESEARCH:
{uc_context}

Your job is to produce a complete specification that covers sections 5, 6, 9, 11, and 12 of the PRD.

Respond in exactly this structure:

CURRENT STATE SUMMARY:
How does this work today? What manual processes exist?
2-3 sentences based on the SOP above.

SCOPE:
IN SCOPE (list only — no adjectives):
- [specific feature or capability]

OUT OF SCOPE (at least 3 items with reason):
- [item] — [one line reason for exclusion]

FUNCTIONAL USE CASES:
EXACTLY 3 use cases only — one happy path, one edge case, one failure:

Use Case [N]: [name]
Actor: [who — primary user / CS agent / ops team / system]
Precondition: [what must be true before this starts]
User Story: As a [persona], I want [specific action] so that [measurable outcome]
Steps:
  1. [step]
  2. [step]
  3. [step]
Acceptance Criteria:
  - Given [context], when [action], then [expected result]
  - Given [context], when [action], then [expected result]
  - Given [context], when [action], then [expected result]
Failure State AC:
  - Given [error condition], when [action], then [user sees specific message/screen]
Priority: P0 / P1 / P2

NON-FUNCTIONAL USE CASES:
Cover performance, security, accessibility, and reliability:
- [NFR name]: [specific measurable requirement]
  Acceptance: [how to verify this]

METRICS IMPACT (maximum 3 most impacted metrics only):
For each metric:
- Metric name: [name]
- Current baseline: [value or "unknown — check with analytics"]
- Expected impact: [direction and magnitude]
- Owner: [team]

NEW METRICS TO TRACK:
For each new metric this feature requires:
- Metric name: [name]
- Definition: [exactly what is measured]
- Target: [goal]
- Owner: [team]

REPORTS & EXTRACTS IMPACT:
- New reports needed: [list]
- Existing reports affected: [list]
- Reconciliation requirements: [what needs to balance against what]
- Data extract requirements: [downstream teams needing data]

INSTRUMENTATION REQUIREMENTS (maximum 5 key events only):
- Event name: [follow naming convention from reference doc]
- Trigger: [exactly when this fires]
- Properties: [key-value pairs to capture]
- Purpose: [what decision this informs]"""

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        temperature=0.0,
        system=[
            {
                "type": "text",
                "text": f"You are a senior product analyst for HDFC PIXEL Studio writing detailed PRD specifications with complete acceptance criteria for QA teams.\n\nPIXEL context:\n{pixel_context}",
                "cache_control": {"type": "ephemeral"}
            }
        ],
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = message.content[0].text

    result = {
        "agent": "spec",
        "feature_name": input.feature_name,
        "timestamp": datetime.now().isoformat(),
        "sop_used": sop_context.split('\n')[0],
        "response": response_text
    }

    saved_path = save_analysis(input.feature_name + "_spec", result)
    result["saved_to"] = saved_path

    return result

# ============================================================
# PRD GENERATOR FROM FORM
# ============================================================

class PRDFormInput(BaseModel):
    problem_statement: str = ""
    current_solution: str = ""
    competitive_context: str = ""
    rbi_requirements: str = ""
    scope: str = ""
    use_cases: str = ""
    cs_impact: str = ""
    ops_impact: str = ""
    metrics: str = ""
    rollout: str = ""
    reports: str = ""
    instrumentation: str = ""


@app.post("/generate-prd-from-form")
def generate_prd_from_form(input: PRDFormInput):

    pixel_context = load_pixel_context()

    prompt = f"""You are an Associate Director of Products at HDFC PIXEL Studio writing a production-grade PRD.

The PM has filled in research and context across 12 sections. Your job is to synthesise this into a complete, well-structured PRD that a VP of Products would approve without sending back for revision.

Use the PM's inputs as the foundation. Enrich where helpful. Never contradict what the PM has provided. If a section is empty, write "To be completed by PM" — do not invent content.

INPUT FROM PM:

1. PROBLEM STATEMENT:
{input.problem_statement or "Not provided"}

2. HOW BANK CURRENTLY SOLVES THIS:
{input.current_solution or "Not provided"}

3. COMPETITIVE CONTEXT:
{input.competitive_context or "Not provided"}

4. RBI & REGULATORY REQUIREMENTS:
{input.rbi_requirements or "Not provided"}

5. SCOPE:
{input.scope or "Not provided"}

6. FUNCTIONAL USE CASES & ACCEPTANCE CRITERIA:
{input.use_cases or "Not provided"}

7. CUSTOMER SUPPORT IMPACT:
{input.cs_impact or "Not provided"}

8. OPERATIONS IMPACT:
{input.ops_impact or "Not provided"}

9. SUCCESS METRICS:
{input.metrics or "Not provided"}

10. ROLLOUT STRATEGY:
{input.rollout or "Not provided"}

11. REPORTS & RECONCILIATION:
{input.reports or "Not provided"}

12. UX & INSTRUMENTATION:
{input.instrumentation or "Not provided"}

Now generate the complete PRD. Follow this exact structure:

---
PIXEL STUDIO — PRODUCT REQUIREMENTS DOCUMENT
---

1. STRATEGIC CONTEXT
Why now? Opportunity cost of not building this in next 2 quarters.
3 sentences maximum.

2. PROBLEM STATEMENT
Rewrite the PM's input in clean PRD language.
User pain + current behaviour + business consequence.
3 sentences maximum.

3. CURRENT STATE
How does this work today? What manual processes exist?

4. COMPETITIVE CONTEXT & DIFFERENTIATION
Synthesise the competitive research. What is PIXEL's chosen angle and why?

5. REGULATORY REQUIREMENTS
List all RBI requirements verbatim. Flag mandatory vs advisory.

6. SCOPE
In scope — feature list only, no adjectives.
Out of scope — at least 3 items with one-line exclusion reason.

7. USER STORIES & ACCEPTANCE CRITERIA
Reproduce and clean up the PM's use cases.
Ensure every story has:
- Given/When/Then ACs
- One failure state AC
- Priority (P0/P1/P2)

8. SUCCESS METRICS
For each metric: name, baseline, target, timeframe, owner.
Split: user metrics vs bank metrics.

9. ROLLOUT STRATEGY
MVP → V1 → V2.
For each phase: what ships, rollout %, kill switch condition, success gate.

10. CUSTOMER SUPPORT & OPERATIONS
CS: new ticket types, SLAs, training needs, deflection strategy.
Ops: manual touchpoints eliminated, new processes required.

11. REPORTS & RECONCILIATION
New reports, existing reports affected, reconciliation requirements, data extracts.

12. UX & INSTRUMENTATION
Key user flows. Events to instrument. Properties to capture.

13. OPEN QUESTIONS
Exactly 3. Format: Question → Why it matters → Who owns the answer.

Be specific to Indian fintech, UPI, and RBI regulatory context.
No generic platitudes. No filler."""

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        temperature=0.0,
        system=f"You are an Associate Director of Products at HDFC PIXEL Studio writing production-grade PRDs for Indian fintech credit card products.\n\n{pixel_context}",
        messages=[{"role": "user", "content": prompt}]
    )

    result = {
        "prd": message.content[0].text,
        "timestamp": datetime.now().isoformat()
    }

    saved_path = save_analysis("prd_generated", result)
    result["saved_to"] = saved_path

    return result


@app.post("/download-prd-docx")
def download_prd_docx(data: dict):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from io import BytesIO
    from fastapi.responses import StreamingResponse

    prd_text = data.get("prd_text", "")

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    title = doc.add_heading('PIXEL Studio — Product Requirements Document', 0)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    for line in prd_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph('')
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename=PIXEL_PRD.docx'}
    )